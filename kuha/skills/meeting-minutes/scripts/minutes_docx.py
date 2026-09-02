#!/usr/bin/env python
"""
Convert a filled-in "Bien ban hop" Markdown file (based on
templates/bien-ban-hop.md) into a .docx document: Times New Roman 13pt,
Vietnamese headings, and a real table for the assignment / signature sections.

Usage:
    py -3.12 minutes_docx.py <markdown-file> [--out file.docx]
"""
import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(13)


def add_run(paragraph, text, bold=False, size=FONT_SIZE):
    run = paragraph.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = size
    run.font.bold = bold
    return run


def add_inline_markdown(paragraph, text, base_size=FONT_SIZE):
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            add_run(paragraph, part[2:-2], bold=True, size=base_size)
        else:
            add_run(paragraph, part, bold=False, size=base_size)


def parse_table_block(lines, start_idx):
    rows = []
    i = start_idx
    while i < len(lines) and lines[i].strip().startswith("|"):
        cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        rows.append(cells)
        i += 1
    if len(rows) >= 2 and re.match(r"^:?-+:?$", rows[1][0].replace(" ", "")):
        rows.pop(1)
    return rows, i


def build_docx(md_path: Path, out_path: Path):
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = FONT_SIZE

    section = doc.sections[0]
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("|"):
            rows, next_i = parse_table_block(lines, i)
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = "Table Grid"
                for r, row_cells in enumerate(rows):
                    for c, cell_text in enumerate(row_cells):
                        if c >= len(table.columns):
                            continue
                        cell = table.cell(r, c)
                        cell.text = ""
                        p = cell.paragraphs[0]
                        add_inline_markdown(p, cell_text, base_size=Pt(12))
                        if r == 0:
                            for run in p.runs:
                                run.font.bold = True
                doc.add_paragraph()
            i = next_i
            continue

        if stripped.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, stripped[2:].strip(), bold=True, size=Pt(16))
            i += 1
            continue

        if stripped.startswith("## "):
            p = doc.add_paragraph()
            add_run(p, stripped[3:].strip(), bold=True, size=Pt(14))
            i += 1
            continue

        if stripped.startswith("### "):
            p = doc.add_paragraph()
            add_run(p, stripped[4:].strip(), bold=True, size=Pt(13))
            i += 1
            continue

        if stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_markdown(p, stripped[2:].strip())
            i += 1
            continue

        if re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            add_inline_markdown(p, re.sub(r"^\d+\.\s", "", stripped))
            i += 1
            continue

        p = doc.add_paragraph()
        add_inline_markdown(p, stripped)
        i += 1

    doc.save(out_path)
    return out_path


def main():
    parser = argparse.ArgumentParser(
        description="Chuyen file Markdown bien ban hop (templates/bien-ban-hop.md da dien) sang .docx"
    )
    parser.add_argument("markdown", help="File .md da dien noi dung")
    parser.add_argument("--out", default=None, help="File .docx output. Mac dinh: cung ten, duoi .docx")
    args = parser.parse_args()

    md_path = Path(args.markdown).resolve()
    if not md_path.is_file():
        print(f"Loi: khong tim thay file: {md_path}", file=sys.stderr)
        sys.exit(1)

    out_path = Path(args.out).resolve() if args.out else md_path.with_suffix(".docx")
    build_docx(md_path, out_path)
    print(f"Da tao file: {out_path}")


if __name__ == "__main__":
    main()
